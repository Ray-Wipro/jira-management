import requests
import sys
from docx import Document
from docx.shared import Pt

# === CONFIG ===
JIRA_URL = "https://ejlog.atlassian.net"        # <-- Modifica con il tuo dominio
USERNAME = "rraimondi@ferrettogroup.com"        # <-- Tuo utente Atlassian
API_TOKEN = "ATATT3xFfGF0rFdHpQs25sHiT1VUn4-jmZl3nsdGWBB_hHAUMa1rePZsj758F_s7d0lWi7C0tCD4a-e8EOTQTvFZO_jtl3WYU0gG5tNv8bg2RkUDV24oYBapEB7-MQF7opXFmFREB1UCMDqRN_-ZZAkQL8xFceDjOowuNTx53CSyB1-HO1WJhUc=9EAB0E42"                  # <-- Generato da https://id.atlassian.com/manage/api-tokens

# === ID dei campi custom ===
CAMPO_RIFERIMENTI = "customfield_10146"
CAMPO_AMBIENTE = "environment"

# === Leggi codice progetto da input ===
if len(sys.argv) < 2:
    print("❌ Inserire il codice progetto (es: DNT-3)")
    sys.exit(1)

CODICE_PROGETTO = sys.argv[1].upper()
PROJECT_KEY = CODICE_PROGETTO.split("-")[0]

HEADERS = {"Accept": "application/json"}
AUTH = (USERNAME, API_TOKEN)

# === Recupero dati issue principale ===
issue_url = f"{JIRA_URL}/rest/api/3/issue/{CODICE_PROGETTO}"
fields = f"summary,{CAMPO_RIFERIMENTI},{CAMPO_AMBIENTE}"

resp = requests.get(issue_url, headers=HEADERS, auth=AUTH)

def estrai_testo_da_campo_rich_text(campo):
    if not campo:
        return "N/D"
    try:
        testo = ""
        for block in campo.get("content", []):
            for inner in block.get("content", []):
                if inner.get("type") == "text":
                    testo += inner.get("text", "")
                elif inner.get("type") == "hardBreak":
                    testo += "\r\n"  # o solo "\n" se preferisci
            testo += "\n"  # fine paragrafo
        return testo.strip()
    except Exception:
        return "N/D"
    
if resp.status_code != 200:
    print(f"❌ Errore nel recupero dell'issue {CODICE_PROGETTO}")
    print(resp.text)
    sys.exit(1)

issue_data = resp.json()
summary = issue_data["fields"]["summary"]
#riferimenti = issue_data["fields"].get(CAMPO_RIFERIMENTI, "N/D")
#ambiente = issue_data["fields"].get(CAMPO_AMBIENTE, "N/D")
riferimenti_raw = issue_data["fields"].get(CAMPO_RIFERIMENTI, None)
riferimenti = estrai_testo_da_campo_rich_text(riferimenti_raw)

ambiente_raw = issue_data["fields"].get(CAMPO_AMBIENTE, None)
ambiente = estrai_testo_da_campo_rich_text(ambiente_raw)

# === Recupero tutte le issue del progetto ===
jql = f'project = "{PROJECT_KEY}" ORDER BY key ASC'
search_url = f"{JIRA_URL}/rest/api/3/search"
params = {
    "jql": jql,
    "fields": "summary",
    "maxResults": 200
}

resp_all = requests.get(search_url, headers=HEADERS, params=params, auth=AUTH)
all_issues = resp_all.json().get("issues", [])

# === Creazione documento Word ===
doc = Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    
# Intestazione
add_heading(f"Documento Progetto - {CODICE_PROGETTO}", level=0)

# Codice progetto e descrizione
add_heading("Descrizione")
add_paragraph(f"{summary}")

# Riferimenti
add_heading("Riferimenti delle persone del cliente")
add_paragraph(riferimenti)

# Ambiente
add_heading("Informazioni sull'Ambiente")
add_paragraph(ambiente)

# Elenco Issues
add_heading("Elenco Issues del Progetto")
for issue in all_issues:
    key = issue["key"]
    title = issue["fields"]["summary"]
    doc.add_paragraph(f"{key} - {title}", style="List Bullet")

# Salvataggio
output_filename = f"{CODICE_PROGETTO}_report.docx"
doc.save(output_filename)

print(f"✅ Documento generato: {output_filename}")
