"""
Script Python per estrarre e documentare i dettagli di un progetto Jira.

Funzionalità principali:
- Recupera tramite API Jira le informazioni di un'issue principale (identificata da codice progetto, es. DNT-3).
- Estrae: titolo, riferimenti del cliente, informazioni sull'ambiente e l'elenco di tutte le issues collegate.
- Analizza e raccoglie tutti i commenti delle issues del progetto, ordinandoli per data.
- Genera un documento Word (.docx) strutturato con:
    - Intestazione progetto
    - Descrizione
    - Riferimenti cliente
    - Ambiente
    - Elenco issues
    - Commenti dettagliati
- Salva il report in un file `<CODICE_PROGETTO>_report.docx`.

Requisiti:
- Librerie Python: requests, python-docx
- API token Atlassian valido
- Permessi di accesso in lettura al progetto Jira

Nome del file:
- jira-project-report-v4.3.py

Autore: Roberto Raimondi
Ultima modifica: 21/08/2025
"""

import os
import requests
import sys

from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm, Pt, RGBColor
from tkinter import Button, Entry, Label, StringVar, Tk, messagebox
from tkinter.ttk import Combobox

# === Caricamento variabili ambiente ===
load_dotenv()

# === CONFIGURAZIONE JIRA ===
VERSION     = "4.3"

JIRA_URL    = os.getenv("JIRA_URL")
USERNAME    = os.getenv("JIRA_USERNAME")
API_TOKEN   = os.getenv("JIRA_API_TOKEN")

AUTH        = (USERNAME, API_TOKEN)
HEADERS     = {"Accept": "application/json"}
JQL_BASE    = 'assignee = currentUser() AND status in ("Da Gestire", "In corso", "Stand by Cliente", "Stand by Interno") ORDER BY key ASC'

# === Costanti per i campi personalizzati ===
CAMPO_AMBIENTE      = "environment"
CAMPO_RIFERIMENTI   = "customfield_10059"

# === Funzione per formattare la data in formato leggibile ===
def _parse_jira_dt(s: str) -> datetime:
    # Jira: "2025-08-13T09:41:22.123+0200" → consideriamo solo la parte fino ai secondi
    return datetime.strptime(s[:19], "%Y-%m-%dT%H:%M:%S")

# === Funzione per ottenere tutti i ticket dell'utente ===
def get_tickets_for_user():
    search_url = f"{JIRA_URL}/rest/api/3/search"
    params = {
        "jql": JQL_BASE,
        "fields": "key, summary",
        "maxResults": 1000
    }
    resp = requests.get(search_url, headers=HEADERS, params=params, auth=AUTH)
    if resp.status_code != 200:
        return []
    
    issues = resp.json().get("issues", [])
    tickets = [f"{i['key']} - {i['fields']['summary']}" for i in issues]
    return sorted(tickets, key=lambda x: x.split(" - ")[0])

# === Estrae tutti i commenti di un progetto ordinati per data crescente (dal più vecchio al più recente)===
def get_ticket_comments(ticket_key):
    url = f"{JIRA_URL}/rest/api/3/issue/{ticket_key}/comment"
    start_at, max_results = 0, 100
    all_comments = []

    while True:
        params = {"startAt": start_at, "maxResults": max_results}
        resp = requests.get(url, headers=HEADERS, params=params, auth=AUTH)
        if resp.status_code != 200:
            break

        data = resp.json()
        comments = data.get("comments", [])
        if not comments:
            break

        for c in comments:
            created = _parse_jira_dt(c.get("created", ""))
            author = (c.get("author") or {}).get("displayName", "Sconosciuto")
            body = c.get("body", None)
            all_comments.append({
                "created": created,
                "author": author,
                "body": body
            })

        start_at += len(comments)
        if start_at >= data.get("total", start_at):
            break

    all_comments.sort(key=lambda x: x["created"])
    return all_comments

# === Funzione per applicare gli stili a un run di testo ===
def apply_marks_to_run(run, marks: list):
    for mark in marks:
        mtype = mark.get("type")
        attrs = mark.get("attrs", {})

        match mtype:
            case "strong":
                run.bold = True
            case "em":
                run.italic = True
            case "underline":
                run.underline = True
            case "strike":
                run.font.strike = True
            case "subsup":
                match attrs.get("subscript"), attrs.get("superscript"):
                    case True, _:
                        run.font.subscript = True
                    case _, True:
                        run.font.superscript = True
            case "color" | "textColor":
                color = attrs.get("color", "000000")
                run.font.color.rgb = RGBColor.from_string(normalize_color(color))
            case "link":
                # Hyperlink non nativo in docx, lasciamo solo il testo visibile
                pass
            case "code":
                run.font.name = "Courier New"
                run.font.size = Pt(9)

# === Funzione per normalizzare i colori esadecimali ===
def normalize_color(color: str) -> str:
    if not color:
        return "000000"  # default nero
    color = color.strip()
    if color.startswith("#"):
        color = color[1:]
    match len(color):
        case 1:
            color = color * 6
        case 3:
            color = "".join([c*2 for c in color])
        case 4:
            # rgba esadecimale tipo #f00f -> ignoriamo alpha
            color = "".join([c*2 for c in color[:3]])
        case _ if len(color) > 6:
            color = color[:6]
    return color.upper()

# === Funzione per aggiungere testo con stili a un paragrafo o cella ===
def add_text(parent, text, marks=None):
    if not text:
        return None
    if marks is None:
        marks = []

    # Determina dove aggiungere il run
    match parent:
        case _Cell():
                p = parent.add_paragraph()
                run = p.add_run(text)
        case Paragraph():
                run = parent.add_run(text)
        case "Document":
            p = parent.add_paragraph()
            run = p.add_run(text)
        case _:
                p = parent.add_paragraph()
                run = p.add_run(text)

    apply_marks_to_run(run, marks)
    return run

# === Funzione per generare il documento Word ===
def parse_adf_to_docx(content, parent, level=1):
    """
    Converte il contenuto ADF (Atlassian Document Format) in paragrafi e run di Word.
    Gestisce paragrafi, hardBreak, grassetto e corsivo.
    """
    for node in content:
        node_type = node.get("type")

        match node_type:
            case "paragraph":
                p = parent.add_paragraph()
                for child in node.get("content", []):
                    child_type = child.get("type")
                    match child_type:
                        case "text":
                            add_text(p, child.get("text", ""), marks=child.get("marks", []))
                        case "hardBreak":
                            p.add_run().add_break()
                        case _ if "content" in child:
                            parse_adf_to_docx(child["content"], p, level)

            case "heading":
                level = node.get("attrs", {}).get("level", 1)
                if "content" in node:
                    heading_text = "".join(
                        [c.get("text", "") for c in node["content"] if c["type"] == "text"]
                    )
                    match type(parent).__name__:
                        case "Document" | "Paragraph":
                            parent.add_heading(heading_text, level=level)
                        case "Cell":
                            # dentro una cella, aggiungi paragrafo con stile heading
                            p = parent.add_paragraph(heading_text, style=f"Heading {level}")

            # case "bulletList" | "orderedList":
            #     style = "List Bullet" if node_type == "bulletList" else "List Number"
            #     for li in node.get("content", []):  # ogni listItem
            #         if li.get("type") == "listItem":
            #             for paragraph_node in li.get("content", []):
            #                 if paragraph_node["type"] == "paragraph":
            #                     p = parent.add_paragraph(style=style)
            #                     for child in paragraph_node.get("content", []):
            #                         child_type = child.get("type")
            #                         match child_type:
            #                             case "text":
            #                                 add_text(p, child.get("text", ""), marks=child.get("marks", []))
            #                             case "hardBreak":
            #                                 p.add_run().add_break()
            #                             case _ if "content" in child:
            #                                 parse_adf_to_docx(child["content"], p, level)
            #                 elif paragraph_node["type"] in ("bulletList", "orderedList"):
            #                     # Gestione corretta delle liste annidate
            #                     parse_adf_to_docx([paragraph_node], parent, level + 1)
                                
            case "bulletList" | "orderedList":
                ordered = node_type == "orderedList"
                for li in node.get("content", []):  # ogni listItem
                    if li.get("type") == "listItem":
                        # 1) Estrae tutto il testo dai paragraph interni
                        text_parts = []
                        for child in li.get("content", []):
                            if child.get("type") == "paragraph":
                                text_parts.append(get_text_from_content(child.get("content", [])))
                        raw_text = "\n".join([t for t in text_parts if t.strip()])
                        if raw_text:
                            add_bullet(parent, raw_text, level)
                        # 2) processa eventuali sotto-liste annidate
                        for child in li.get("content", []):
                            if child.get("type") in ("bulletList", "orderedList"):
                                parse_adf_to_docx([child], parent, level + 1)
                    
            case "codeBlock":
                code_text = ""
                for child in node.get("content", []):
                    if child["type"] == "text":
                        code_text += child.get("text", "") + "\n"
                if code_text.strip():
                    p = parent.add_paragraph()
                    run = add_text(p, code_text.rstrip(), marks=node.get("marks", []))
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)

            case "panel":
                table = parent.add_table(rows=1, cols=1)
                cell = table.rows[0].cells[0]
                add_info_panel(cell)
                if "content" in node:
                    parse_adf_to_docx(node["content"], cell, level)

            case "text":
                text = node.get("text", "")
                marks = node.get("marks", [])
                add_text(parent, text, marks)

            case "hardBreak":
                if isinstance(parent, Paragraph):
                    parent.add_run().add_break()
                elif isinstance(parent, _Cell):
                    parent.add_paragraph("")
                else:
                    parent.add_paragraph("")

            case _ if "content" in node:
                parse_adf_to_docx(node["content"], parent, level)

            case _:
                pass
        # Se ci sono altri tipi di nodo, si possono aggiungere qui...

# === Funzione per aggiungere un pannello informativo con bordo e sfondo ===
def add_info_panel(cell, bg_color="D9D9D9", border_size=4, border_color="000000"):
    """
    Trasforma una cella (panel) in un pannello visivamente evidenziato:
    - Colore di sfondo (bg_color, esadecimale senza '#')
    - Bordo attorno alla cella (border_size in punti, border_color esadecimale)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Shading (sfondo)
    shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
    tcPr.append(shading)

    # Bordo (perimetrale)
    from docx.oxml import OxmlElement
    borders = OxmlElement('w:tcBorders')
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        borders.append(border)
    tcPr.append(borders)

    # Aggiunta di padding interno (optional)
    tcPr.set(qn('w:textDirection'), "btLr")

# === Funzione per gestire elenchi puntati e numerati con indentazione manuale ===
def parse_list(node, parent, level=1, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    for li in node.get("content", []):
        # crea il paragrafo principale del bullet
        p = parent.add_paragraph(style=style)
        p.paragraph_format.left_indent = Cm(0.75 * (level - 1))
        p.paragraph_format.first_line_indent = Cm(0)

        for child in li.get("content", []):
            match child.get("type"):
                case "paragraph":
                    for sub in child.get("content", []):
                        match sub.get("type"):
                            case "text":
                                add_text(p, sub.get("text", ""), marks=sub.get("marks", []))
                            case "hardBreak":
                                p.add_run().add_break()
                            case _ if "content" in sub:
                                # testo annidato dentro paragraph
                                parse_adf_to_docx(sub["content"], p, level + 1)
                case "bulletList":
                    # sotto-elenco: livello +1
                    parse_list(child, parent, level + 1, ordered=False)
                case "orderedList":
                    parse_list(child, parent, level + 1, ordered=True)

# === Funzione per estrarre il testo da un contenuto ADF (rich text) ===
def get_text_from_content(content_list):
    texts = []
    for c in content_list:
        match c.get("type"):
            case "text":
                texts.append(c.get("text", ""))

            case "paragraph":
                if "content" in c:
                    texts.append(get_text_from_content(c["content"]))
                texts.append("\n")  # newline dopo ogni paragrafo

            case "hardBreak":
                texts.append("\n")

            case "bulletList" | "orderedList":
                for li in c.get("content", []):
                    texts.append(get_text_from_content(li.get("content", [])))

            case "listItem":
                if "content" in c:
                    texts.append(get_text_from_content(c["content"]))
                texts.append("\n")  # newline dopo ogni listItem

            case "panel":
                if "content" in c:
                    texts.append(get_text_from_content(c["content"]))

            case _:
                if "content" in c:
                    texts.append(get_text_from_content(c["content"]))

    return "".join(texts)

# === Estrazione commenti da issues del progetto ===
def parse_rich_text(raw_field):
    if not raw_field:
        return ""
    if isinstance(raw_field, dict) and "content" in raw_field:
        return get_text_from_content(raw_field["content"])
    if isinstance(raw_field, str):
        return raw_field
    return ""

# === Recupero dettagli ticket ===
def get_ticket_details(ticket_key):
    url = f"{JIRA_URL}/rest/api/3/issue/{ticket_key}"
    params = {"fields": f"summary, description, {CAMPO_RIFERIMENTI}, {CAMPO_AMBIENTE}, project"}
    resp = requests.get(url, headers=HEADERS, params=params, auth=AUTH)
    if resp.status_code != 200:
        return None
    
    fields = resp.json()["fields"]
    
    summary = fields.get("summary", "")
    description = fields.get("description", {})
    riferimenti = parse_rich_text(fields.get(CAMPO_RIFERIMENTI, {}))
    ambiente = parse_rich_text(fields.get(CAMPO_AMBIENTE, {}))
    cliente = fields.get("project", {}).get("name", "-")

    return summary, description, riferimenti, ambiente, cliente

# === Aggiunta testo multilinea in Word ===
def add_multiline_text(parent, text: str):
    """Aggiunge testo multilinea (con \n) rispettando gli a capo in Word."""
    if not text:
        return
    for line in text.splitlines():
        p = parent.add_paragraph(line.strip() if line.strip() else "")
    return

# === Funzione per aggiungere un bullet con indentazione coerente ===
def add_bullet(doc, text, level=1):
    if not text:
        return
    lines = text.splitlines()
    first_line = lines[0] if lines else ""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.left_indent = Cm(0.75 * (level - 1))
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(2)
    para.add_run("• " + first_line)

    # righe successive del bullet
    for extra_line in lines[1:]:
        p2 = doc.add_paragraph()
        pf2 = p2.paragraph_format
        pf2.left_indent = Cm(0.75 * level)
        pf2.first_line_indent = Cm(0)
        pf2.space_after = Pt(2)
        p2.add_run(extra_line)

    return para

# === Sostituire il case "bulletList" | "orderedList" in parse_adf_to_docx ===

# === Creazione documento Word ===
def create_word_document(ticket_key, summary, description_adf, riferimenti, ambiente, comments, cliente):
    doc = Document()

    # Imposta margini pagina
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Definisci stile paragrafo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    # Per applicare Arial correttamente anche a caratteri asiatici ecc.
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Interlinea singola e nessuno spazio tra paragrafi dello stesso stile
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Intestazione
    doc.add_heading(f"{cliente} - {ticket_key}", level=0)

    # Descrizione breve
    doc.add_heading("Descrizione", level=1)
    doc.add_paragraph(summary or "-")

    # Riferimenti
    doc.add_heading("Riferimenti delle persone del cliente", level=1)
    if isinstance(riferimenti, list):
        for ref in riferimenti:
            add_bullet(doc, ref)
    elif isinstance(riferimenti, str):
        add_multiline_text(doc, riferimenti.strip())
    else:
        doc.add_paragraph("-")

    # Ambiente
    doc.add_heading("Informazioni sull'Ambiente", level=1)
    add_multiline_text(doc, ambiente)

    # Descrizione dettagliata
    doc.add_heading("Descrizione dettagliata", level=1)
    if isinstance(description_adf, dict) and "content" in description_adf:
        parse_adf_to_docx(description_adf["content"], doc)
    elif isinstance(description_adf, str):
        doc.add_paragraph(description_adf.strip())
    else:
        doc.add_paragraph("(Nessuna descrizione fornita)")

    # Elenco Commenti
    doc.add_heading("Commenti del Progetto", level=1)
    if comments:
        for c in comments:
            header = f"[{c['created'].strftime('%d-%m-%Y %H:%M')}] {c['author']}"
            p = doc.add_paragraph()
            # run = p.add_run(header)
            run = add_text(p, header, marks=[{"type": "strong"}])
            run.bold = True

            body = c["body"]
            if isinstance(body, dict) and "content" in body:
                parse_adf_to_docx(body["content"], doc)
            elif isinstance(body, str):
                add_multiline_text(doc, body.strip())
            else:
                doc.add_paragraph("—")

            doc.add_paragraph("")  # spazio tra commenti
    else:
        doc.add_paragraph("(Nessun commento)")

    # Salvataggio file
    filename = f"{ticket_key}_report.docx"
    doc.save(filename)
    print(f"Documento salvato: {filename}")

# === GUI selezione ticket ===
def select_ticket_gui(tickets_list):
    root = Tk()
    root.title("Selezione Ticket Jira")
    root.geometry("500x200")

    Label(root, text="Seleziona un ticket aperto:").pack(padx=10, pady=5)
    selected_ticket = StringVar()
    combo = Combobox(root, values=tickets_list, textvariable=selected_ticket, state="readonly", width=60)
    combo.pack(padx=10, pady=5)

    Label(root, text="Oppure inserisci codice ticket (es. XXX-123):").pack(padx=10, pady=5)
    manual_ticket = Entry(root, width=30)
    manual_ticket.pack(padx=10, pady=5)
    manual_ticket.focus_set()

    def on_confirm():
        choice = manual_ticket.get().strip()
        if choice:
            root.selected = choice
        elif selected_ticket.get():
            root.selected = selected_ticket.get().split(" - ")[0]
        else:
            messagebox.showerror("Errore", "Seleziona un ticket o inseriscine uno manualmente.")
            return
        root.destroy()

    Button(root, text="Conferma", command=on_confirm).pack(pady=10)
    manual_ticket.bind("<Return>", lambda event: on_confirm())
    combo.bind("<Return>", lambda event: on_confirm())

    root.mainloop()
    return getattr(root, "selected", None)

# === Main ===
if __name__ == "__main__":
    if len(sys.argv) > 1:
        ticket_key = sys.argv[1]
    else:
        tickets = get_tickets_for_user()
        ticket_key = select_ticket_gui(tickets)

    if not ticket_key:
        print("Nessun ticket selezionato.")
        sys.exit(1)

    print(f"Recupero dettagli per {ticket_key}...")
    details = get_ticket_details(ticket_key)
    if not details:
        print("Errore nel recupero ticket.")
        sys.exit(1)

    summary, description_adf, riferimenti, ambiente, cliente = details
    comments = get_ticket_comments(ticket_key)
    create_word_document(ticket_key, summary, description_adf, riferimenti, ambiente, comments, cliente)

# === Fine script ===

# === Roadmap future ===
# - Aggiungere Stile documento
"""
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
"""
# - Uso della GUI
"""
Attualmente se l’utente inserisce un codice manuale, non viene validato: potresti verificare che rispetti il formato XXX-123 prima di accettarlo.
Potresti anche pre-selezionare il primo ticket nella combobox per velocizzare.
"""
