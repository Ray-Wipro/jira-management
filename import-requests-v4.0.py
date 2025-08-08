import requests
import sys
import os
from datetime import datetime
from docx import Document
from dotenv import load_dotenv

doc = Document()

# === CONFIGURAZIONE ===
JIRA_URL = "https://ejlog.atlassian.net"
USERNAME = "rraimondi@ferrettogroup.com"
load_dotenv()                                   # Carica le variabili d'ambiente da .env se presente
API_TOKEN   = os.getenv("JIRA_API_TOKEN")       # <-- Imposta la variabile d'ambiente JIRA_API_TOKEN con il tuo token API

JQL = 'assignee = currentUser() AND status in ("Da Gestire", "In corso", "Stand by Cliente", "Stand by Interno") ORDER BY priority DESC, project, duedate ASC, created ASC'

# === PARAMETRI RICHIESTA ===
url = f"{JIRA_URL}/rest/api/3/search"
headers = {"Accept": "application/json"}
auth = (USERNAME, API_TOKEN)

params = {
    "jql": JQL,
    "fields": "summary,status,priority,created,duedate,project,key",
    "maxResults": 1000
}

# === ESECUZIONE ===
response = requests.get(url, headers=headers, params=params, auth=auth)

if response.status_code != 200:
    print("Errore:", response.status_code)
    print(response.text)
    exit()

issues = response.json().get("issues", [])
print(f"Trovate {len(issues)} issue assegnate")

# === CATEGORIZZAZIONE PER PRIORITÀ ===
priorities = {
    "High": [],
    "Medium": [],
    "Low": [],
    "Nessuna": []
}

def parse_date(date_str):
    return datetime.strptime(date_str, "%Y-%m-%d") if date_str else None

def parse_created(date_str):
    return datetime.strptime(date_str, "%Y-%m-%d")

for issue in issues:
    fields = issue["fields"]
    key = issue["key"]
    project_name = fields["project"]["name"]
    title = fields["summary"]
    status = fields["status"]["name"]
    priority = fields.get("priority", {}).get("name", "Nessuna")
    duedate = fields.get("duedate")  # può essere None
    created = fields["created"][:10]

    title_clean = title
    if title.lower().startswith(project_name.lower()):
        title_clean = title[len(project_name):].lstrip(" -:–—")

    priorities.setdefault(priority, []).append({
        "key": key,
        "progetto": project_name,
        "titolo": title_clean,
        "stato": status,
        "scadenza": duedate,
        "creazione": created
    })

# === ORDINAMENTO E GENERAZIONE OUTPUT ===
output_lines = []

for prio_label in ["High", "Medium", "Low", "Nessuna"]:
    blocco = priorities.get(prio_label, [])
    if not blocco:
        continue

    # Sezione intestazione
    doc.add_paragraph("##############################")
    doc.add_paragraph(f"# {prio_label.upper()} PRIORITY")
    doc.add_paragraph("##############################\n")

    output_lines.append("##############################")
    output_lines.append(f"# {prio_label.upper()} PRIORITY")
    output_lines.append("##############################\n")

    # Ordinamento: prima per scadenza, poi per creazione
    blocco.sort(key=lambda x: (
        parse_date(x["scadenza"]) if x["scadenza"] else datetime.max,
        parse_created(x["creazione"])
    ))

    for item in blocco:
        scad = f", scad. {datetime.strptime(item['scadenza'], '%Y-%m-%d').strftime('%d-%m-%Y')}" if item["scadenza"] else ""
        line = f"{item['progetto']} - {item['titolo']} ({item['stato']}{scad})"

        # Word: key in grassetto
        p = doc.add_paragraph()
        run_key = p.add_run(f"{item['key']} ")
        run_key.bold = True
        p.add_run(line)

        # TXT: key inclusa
        output_lines.append(f"{item['key']} {line}")

    doc.add_paragraph("")
    output_lines.append("")

# === SALVA I FILE ===
doc.save("elenco_attivita.docx")

with open("elenco_attivita.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(output_lines))

print("✅ File 'elenco_attivita.docx' e 'elenco_attivita.txt' generati correttamente.")
