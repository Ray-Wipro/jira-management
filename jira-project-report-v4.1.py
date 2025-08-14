"""
Script Python per estrarre e documentare i dettagli di un progetto Jira.

Funzionalità principali:
- Recupera tramite API Jira le informazioni di un'issue principale (identificata da codice progetto, es. DNT-3).
- Estrae: titolo, riferimenti del cliente, informazioni sull'ambiente e l'elenco di tutte le issues collegate.
- Analizza e raccoglie tutti i commenti delle issues del progetto, ordinandoli per data.
- Genera un documento Word (.docx) strutturato con:
    - Intestazione progetto
    - Descrizione
    - Riferimenti cliente
    - Ambiente
    - Elenco issues
    - Commenti dettagliati
- Salva il report in un file `<CODICE_PROGETTO>_report.docx`.

Requisiti:
- Librerie Python: requests, python-docx
- API token Atlassian valido
- Permessi di accesso in lettura al progetto Jira

Nome del file:
- jira-project-report-v4.1.py

Autore: Roberto Raimondi
Ultima modifica: 14/08/2025
"""

import requests
import sys
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from dotenv import load_dotenv
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from tkinter import Entry, Tk, Label, Button, StringVar, messagebox
from tkinter.ttk import Combobox

# === Caricamento variabili ambiente ===
load_dotenv()

# === CONFIGURAZIONE JIRA ===
VERSIONE    = "4.1"
JIRA_URL    = os.getenv("JIRA_URL")
USERNAME    = os.getenv("JIRA_USERNAME")
API_TOKEN   = os.getenv("JIRA_API_TOKEN")
AUTH        = (USERNAME, API_TOKEN)
HEADERS     = {"Accept": "application/json"}
JQL_BASE = 'assignee = currentUser() AND status in ("Da Gestire", "In corso", "Stand by Cliente", "Stand by Interno") ORDER BY key ASC'

CAMPO_RIFERIMENTI = "customfield_10059"
CAMPO_AMBIENTE = "environment"

# === Funzione per formattare la data in formato leggibile ===
def _parse_jira_dt(s: str) -> datetime:
    # Jira: "2025-08-13T09:41:22.123+0200" → consideriamo solo la parte fino ai secondi
    return datetime.strptime(s[:19], "%Y-%m-%dT%H:%M:%S")

# === Funzione per ottenere tutti i ticket dell'utente ===
def get_tickets_for_user():
    search_url = f"{JIRA_URL}/rest/api/3/search"
    params = {
        "jql": JQL_BASE,
        "fields": "key,summary",
        "maxResults": 1000
    }
    resp = requests.get(search_url, headers=HEADERS, params=params, auth=AUTH)
    if resp.status_code != 200:
        return []
    
    issues = resp.json().get("issues", [])
    tickets = [f"{i['key']} - {i['fields']['summary']}" for i in issues]
    return sorted(tickets, key=lambda x: x.split(" - ")[0])

# === Estrae tutti i commenti di un progetto ordinati per data crescente (dal più vecchio al più recente)===
def get_ticket_comments(ticket_key):
    url = f"{JIRA_URL}/rest/api/3/issue/{ticket_key}/comment"
    start_at, max_results = 0, 100
    all_comments = []

    while True:
        params = {"startAt": start_at, "maxResults": max_results}
        resp = requests.get(url, headers=HEADERS, params=params, auth=AUTH)
        if resp.status_code != 200:
            break

        data = resp.json()
        comments = data.get("comments", [])
        if not comments:
            break

        for c in comments:
            created = _parse_jira_dt(c.get("created", ""))
            author = (c.get("author") or {}).get("displayName", "Sconosciuto")
            body = c.get("body", None)
            all_comments.append({
                "created": created,
                "author": author,
                "body": body
            })

        start_at += len(comments)
        if start_at >= data.get("total", start_at):
            break

    all_comments.sort(key=lambda x: x["created"])
    return all_comments

# === Funzione per aggiungere un bullet point con indentazione ===
def add_bullet(doc, text, level=0):
    if not text.strip():
        return

    # Seleziona lo stile in base al livello
    style_name = "List Bullet" if level == 0 else f"List Bullet {level+1}"
    
    # Crea un paragrafo con lo stile di elenco puntato
    p = doc.add_paragraph(style=style_name)
    
    # Inserisce il testo mantenendo i ritorni a capo interni
    for i, line in enumerate(text.splitlines()):
        if i > 0:
            p.add_run("\n")  # ritorno a capo interno
        p.add_run(line)

# === Funzione di parsing di un nodo ADF in un documento Word con elenchi annidati ===
def parse_node(node, doc, level=0):
    ntype = node.get("type")
    if ntype is None:
        return

    match ntype:
        case "heading":
            lvl = node.get("attrs", {}).get("level", 1)
            text = get_text_from_content(node.get("content", []))
            if text.strip():
                doc.add_heading(text, level=lvl)

        case "paragraph":
            text = get_text_from_content(node.get("content", []))
            if text.strip():
                doc.add_paragraph(text)

        case "bulletList" | "orderedList":
            for li in node.get("content", []):
                parse_node(li, doc, level)

        case "listItem":
            # Testo del listItem (escludendo eventuali sotto-liste)
            text_parts = [
                get_text_from_content([child])
                for child in node.get("content", [])
                if not (isinstance(child, dict) and child.get("type") in ("bulletList", "orderedList"))
            ]
            raw_text = "\n".join([t for t in text_parts if t is not None])

            if raw_text.strip():
                add_bullet(doc, raw_text, level)

            # Sotto-liste (ricorsione con livello+1)
            for child in node.get("content", []):
                if isinstance(child, dict) and child.get("type") in ("bulletList", "orderedList"):
                    parse_node(child, doc, level + 1)
                    
        case "table":
            rows = node.get("content", [])
            if rows:
                cols_count = len(rows[0].get("content", []))
                table = doc.add_table(rows=len(rows), cols=cols_count)
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row.get("content", [])):
                        cell_text = get_text_from_content(cell.get("content", []))
                        table.cell(r_idx, c_idx).text = cell_text

        case _:  # altri tipi non gestiti esplicitamente
            for child in node.get("content", []):
                if isinstance(child, dict):
                    parse_node(child, doc, level)

# === Funzione per generare il documento Word ===
def parse_adf_to_docx(content, doc):
    for block in content:
        parse_node(block, doc)

# === Funzione per estrarre il testo da un contenuto ADF (rich text) ===
def get_text_from_content(content_list):
    texts = []
    for c in content_list:
        match c.get("type"):
            case "text":
                texts.append(c.get("text", ""))
            case _ if "content" in c:
                texts.append(get_text_from_content(c["content"]))
            case _:
                pass
    return "".join(texts)
    
# === Estrazione commenti da issues del progetto ===
def parse_rich_text(raw_field):
    if not raw_field:
        return ""
    if isinstance(raw_field, dict) and "content" in raw_field:
        return get_text_from_content(raw_field["content"])
    if isinstance(raw_field, str):
        return raw_field
    return ""

# === Recupero dettagli ticket ===
def get_ticket_details(ticket_key):
    url = f"{JIRA_URL}/rest/api/3/issue/{ticket_key}"
    params = {"fields": f"summary, description, {CAMPO_RIFERIMENTI}, {CAMPO_AMBIENTE}, project"}
    resp = requests.get(url, headers=HEADERS, params=params, auth=AUTH)
    if resp.status_code != 200:
        return None
    
    fields = resp.json()["fields"]
    
    summary = fields.get("summary", "")
    description = fields.get("description", {})
    riferimenti = parse_rich_text(fields.get(CAMPO_RIFERIMENTI, {}))
    ambiente = parse_rich_text(fields.get(CAMPO_AMBIENTE, {}))
    cliente = fields.get("project", {}).get("name", "-")

    return summary, description, riferimenti, ambiente, cliente

# === Creazione documento Word ===
def create_word_document(ticket_key, summary, description_adf, riferimenti, ambiente, comments, cliente):
    doc = Document()

    # Imposta margini pagina
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Definisci stile paragrafo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    # Per applicare Arial correttamente anche a caratteri asiatici ecc.
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Interlinea singola e nessuno spazio tra paragrafi dello stesso stile
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Intestazione
    doc.add_heading(f"{cliente} - {ticket_key}", level=0)

    # Descrizione breve
    doc.add_heading("Descrizione", level=1)
    doc.add_paragraph(summary or "-")

    # Riferimenti
    doc.add_heading("Riferimenti delle persone del cliente", level=1)
    doc.add_paragraph(riferimenti or "-")

    # Ambiente
    doc.add_heading("Informazioni sull'Ambiente", level=1)
    doc.add_paragraph(ambiente or "-")

    # Descrizione dettagliata
    doc.add_heading("Descrizione dettagliata", level=1)
    if isinstance(description_adf, dict) and "content" in description_adf:
        parse_adf_to_docx(description_adf["content"], doc)
    elif isinstance(description_adf, str):
        doc.add_paragraph(description_adf.strip())
    else:
        doc.add_paragraph("(Nessuna descrizione fornita)")

    # Elenco Commenti
    doc.add_heading("Commenti del Progetto", level=1)
    if comments:
        for c in comments:
            header = f"[{c['created'].strftime('%Y-%m-%d %H:%M')}] {c['author']}"
            p = doc.add_paragraph()
            run = p.add_run(header)
            run.bold = True

            body = c["body"]
            if isinstance(body, dict) and "content" in body:
                parse_adf_to_docx(body["content"], doc)
            elif isinstance(body, str):
                text = body.strip()
                if text:
                    doc.add_paragraph(text)
            else:
                doc.add_paragraph("—")

            doc.add_paragraph("")  # spazio tra commenti
    else:
        doc.add_paragraph("(Nessun commento)")

    # Salvataggio file
    filename = f"{ticket_key}_report.docx"
    doc.save(filename)
    print(f"Documento salvato: {filename}")


# === GUI selezione ticket ===
def select_ticket_gui(tickets_list):
    root = Tk()
    root.title("Selezione Ticket Jira")
    root.geometry("500x200")

    Label(root, text="Seleziona un ticket:").pack(padx=10, pady=5)
    selected_ticket = StringVar()
    combo = Combobox(root, values=tickets_list, textvariable=selected_ticket, state="readonly", width=60)
    combo.pack(padx=10, pady=5)

    Label(root, text="Oppure inserisci codice ticket (es. XXX-123):").pack(padx=10, pady=5)
    manual_ticket = Entry(root, width=30)
    manual_ticket.pack(padx=10, pady=5)

    def on_confirm():
        choice = manual_ticket.get().strip()
        if choice:
            root.selected = choice
        elif selected_ticket.get():
            root.selected = selected_ticket.get().split(" - ")[0]
        else:
            messagebox.showerror("Errore", "Seleziona un ticket o inseriscine uno manualmente.")
            return
        root.destroy()

    Button(root, text="Conferma", command=on_confirm).pack(pady=10)
    root.mainloop()
    return getattr(root, "selected", None)

# === Main ===
if __name__ == "__main__":
    if len(sys.argv) > 1:
        ticket_key = sys.argv[1]
    else:
        tickets = get_tickets_for_user()
        ticket_key = select_ticket_gui(tickets)

    if not ticket_key:
        print("Nessun ticket selezionato.")
        sys.exit(1)

    print(f"Recupero dettagli per {ticket_key}...")
    details = get_ticket_details(ticket_key)
    if not details:
        print("Errore nel recupero ticket.")
        sys.exit(1)

    summary, description_adf, riferimenti, ambiente, cliente = details
    comments = get_ticket_comments(ticket_key)
    create_word_document(ticket_key, summary, description_adf, riferimenti, ambiente, comments, cliente)

# === Fine script ===

# === Roadmap future ===
# - Aggiungere Stile documento
"""
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
"""
# - Uso della GUI
"""
Attualmente se l’utente inserisce un codice manuale, non viene validato: potresti verificare che rispetti il formato XXX-123 prima di accettarlo.
Potresti anche pre-selezionare il primo ticket nella combobox per velocizzare.
"""
