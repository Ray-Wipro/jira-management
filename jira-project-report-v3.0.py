"""
Script Python per estrarre e documentare i dettagli di un progetto Jira.

Funzionalità principali:
- Recupera tramite API Jira le informazioni di un'issue principale (identificata da codice progetto, es. DNT-3).
- Estrae: titolo, riferimenti del cliente, informazioni sull'ambiente e l'elenco di tutte le issues collegate.
- Analizza e raccoglie tutti i commenti delle issues del progetto, ordinandoli per data.
- Genera un documento Word (.docx) strutturato con:
    - Intestazione progetto
    - Descrizione
    - Riferimenti cliente
    - Ambiente
    - Elenco issues
    - Commenti dettagliati
- Salva il report in un file `<CODICE_PROGETTO>_report.docx`.

Requisiti:
- Librerie Python: requests, python-docx
- API token Atlassian valido
- Permessi di accesso in lettura al progetto Jira

Nome del file: 
- jira-project-report-v3.0.py

Autore: Roberto Raimondi
"""

import requests
import sys
import os
from docx import Document
from docx.shared import Pt
from datetime import datetime
from dotenv import load_dotenv
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from tkinter import Entry, Tk, Label, Button, StringVar, messagebox
from tkinter.ttk import Combobox

load_dotenv()

# === CONFIG ===
VERSIONE    = "3.0"
JIRA_URL    = os.getenv("JIRA_URL")
USERNAME    = os.getenv("JIRA_USERNAME")
API_TOKEN   = os.getenv("JIRA_API_TOKEN")

HEADERS = {"Accept": "application/json"}
AUTH = (USERNAME, API_TOKEN)

#CODICE_PROGETTO = sys.argv[1].upper()
#PROJECT_KEY = CODICE_PROGETTO.split("-")[0]

JQL = 'assignee = currentUser() AND status in ("Da Gestire", "In corso", "Stand by Cliente", "Stand by Interno") ORDER BY priority DESC, project, duedate ASC, created ASC'

# JQL = f'project = "{PROJECT_KEY}" ORDER BY created ASC'
# PARAMS = {
#     "jql": JQL,
#     "fields": "key",
#     "maxResults": 1000
# }

# === ID dei campi custom ===
#CAMPO_RIFERIMENTI = "customfield_10146"
CAMPO_RIFERIMENTI = "customfield_10059"
CAMPO_AMBIENTE = "environment"

def estrai_testo_rich_text(campo):
    if not campo or "content" not in campo:
        return ""
    testo = ""
    for blocco in campo["content"]:
        if blocco["type"] == "text":
            testo += blocco.get("text", "")
        elif blocco["type"] == "hardBreak":
            testo += "\n"
        else:
            # ricorsione per contenuti annidati
            testo += estrai_testo_rich_text(blocco)
        # Aggiungo una nuova riga dopo ogni blocco (es. paragrafo)
        if blocco["type"] in ("paragraph", "heading", "blockquote"):
            testo += "\n"
    return testo.strip()

def estrai_commenti_progetto(jira_url, auth, issues, codice_progetto):
    tutti_commenti = []
    for issue in issues:
        if not issue["key"].startswith(codice_progetto):
            continue
        issue_key = issue["key"]
        url_commenti = f"{jira_url}/rest/api/3/issue/{issue_key}/comment"
        resp = requests.get(url_commenti, auth=auth, headers={"Accept": "application/json"})
        if resp.status_code == 200:
            dati_commenti = resp.json().get("comments", [])
            for c in dati_commenti:
                data_commento = datetime.strptime(c["created"][:19], "%Y-%m-%dT%H:%M:%S")
                data_formattata = data_commento.strftime("%d %b %Y")  # '26 set 2023'
                autore = c["author"]["displayName"]

                # Qui elaboriamo il corpo come rich text
                body_field = c.get("body", {})
                testo = estrai_testo_rich_text(c.get("body", {}))

                tutti_commenti.append({
                    "issue_key": issue_key,
                    "created": data_commento,
                    "author": autore,
                    "body": testo
                })
    tutti_commenti.sort(key=lambda x: x["created"])
    return tutti_commenti

def get_all_tickets():
    """Recupera tutti i ticket accessibili a questo utente"""
    search_url = f"{JIRA_URL}/rest/api/3/search"
    params = {
        "jql": JQL,
        "fields": "summary",
        "maxResults": 1000
    }

    resp = requests.get(search_url, headers=HEADERS, params=params, auth=AUTH)
    if resp.status_code == 200:
        issues = resp.json().get("issues", [])
        issues_sorted = sorted(issues, key=lambda i: (i['key'].split('-')[0], int(i['key'].split('-')[1])))
        return [f"{i['key']} - {i['fields']['summary']}" for i in issues_sorted]
    else:
        return []

def avvia_generazione(ticket):
    ticket = ticket.strip().upper()
    if not ticket:
        messagebox.showerror("Errore", "Devi selezionare o inserire un ticket.")
        return
    root.destroy()  # Chiudi la finestra GUI
    genera_documento(ticket.split(" ")[0])

def avvia_gui():
    global root
    root = Tk()
    root.title("Seleziona Ticket Jira")

    Label(root, text="Ticket manuale:").pack(pady=5)
    ticket_var = StringVar()
    Entry(root, textvariable=ticket_var, width=30).pack()

    Label(root, text="Oppure seleziona dalla lista:").pack(pady=5)
    combo_var = StringVar()
    combo = Combobox(root, textvariable=combo_var, width=50)
    combo['values'] = get_all_tickets()
    combo.pack()

    Button(root, text="Conferma", command=lambda: avvia_generazione(
        ticket_var.get() if ticket_var.get().strip() else combo_var.get()
    )).pack(pady=10)

    root.mainloop()

def genera_documento(ticket_code):
    HEADERS = {"Accept": "application/json"}
    AUTH = (USERNAME, API_TOKEN)
    CODICE_PROGETTO = ticket_code
    PROJECT_KEY = CODICE_PROGETTO.split("-")[0]

    # Recupero dati issue principale (quella del codice completo esatto)
    issue_url = f"{JIRA_URL}/rest/api/3/issue/{CODICE_PROGETTO}"
    resp = requests.get(issue_url, headers=HEADERS, auth=AUTH)

    if resp.status_code != 200:
        print(f"❌ Errore nel recupero dell'issue {CODICE_PROGETTO}")
        print(resp.text)
        sys.exit(1)

    issue_data = resp.json()
    summary = issue_data["fields"]["summary"]

    riferimenti_raw = issue_data["fields"].get(CAMPO_RIFERIMENTI, None)
    riferimenti = estrai_testo_rich_text(riferimenti_raw)

    ambiente_raw = issue_data["fields"].get(CAMPO_AMBIENTE, None)
    ambiente = estrai_testo_rich_text(ambiente_raw)

   # Recupero tutte le issue del progetto
    jql = f'project = "{PROJECT_KEY}" ORDER BY created ASC'
    search_url = f"{JIRA_URL}/rest/api/3/search"
    PARAMS = {
        "jql": jql,
        "fields": "summary",
        "maxResults": 1000
    }

    resp_all = requests.get(search_url, headers=HEADERS, params=PARAMS, auth=AUTH)
    all_issues = resp_all.json().get("issues", [])

    # Filtra solo le issue esatte del progetto (es. tutte DNT-3, DNT-4 scartate)
    filtered_issues = [iss for iss in all_issues if iss["key"].startswith(CODICE_PROGETTO)]

    # Estrai commenti solo da queste issue
    commenti = estrai_commenti_progetto(JIRA_URL, AUTH, filtered_issues, CODICE_PROGETTO)

    # Creazione documento Word
    doc = Document()

    # Imposta margini pagina
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)      # margine superiore ridotto (0.5 cm)
        section.bottom_margin = Cm(0.5)   # margine inferiore ridotto (0.5 cm)
        section.left_margin = Cm(1)       # margine sinistro 1 cm
        section.right_margin = Cm(1)      # margine destro 1 cm

    # Definisci stile paragrafo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    # Per applicare Arial correttamente anche a caratteri asiatici ecc.
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Interlinea singola e nessuno spazio tra paragrafi dello stesso stile
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    def add_heading(text, level=1):
        doc.add_heading(text, level=level)

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        run.font.size = Pt(11)

    # Intestazione
    add_heading(f"Documento Progetto - {CODICE_PROGETTO}", level=0)

    # Descrizione
    add_heading("Descrizione")
    add_paragraph(summary)

    # Riferimenti
    add_heading("Riferimenti delle persone del cliente")
    add_paragraph(riferimenti)

    # Ambiente
    add_heading("Informazioni sull'Ambiente")
    add_paragraph(ambiente)

    # Elenco Issues
    add_heading("Elenco Issues del Progetto")
    for issue in filtered_issues:
        key = issue["key"]
        title = issue["fields"]["summary"]
        doc.add_paragraph(f"{key} - {title}", style="List Bullet")

    doc.add_page_break()

    # Commenti
    add_heading("Commenti di tutte le issues del progetto", level=1)
    for c in commenti:
        created_str = c["created"].strftime("%Y-%m-%d %H:%M")
        p = doc.add_paragraph()
        run = p.add_run(f"{c['issue_key']} - {created_str} - {c['author']}")
        run.bold = True
        doc.add_paragraph(c["body"])
        doc.add_paragraph("")

    # Salvataggio
    output_filename = f"{CODICE_PROGETTO}_report.docx"
    doc.save(output_filename)

    print(f"✅ Documento generato: {output_filename}")

# === MAIN ===
if len(sys.argv) > 1:
    genera_documento(sys.argv[1].upper())
else:
    avvia_gui()
