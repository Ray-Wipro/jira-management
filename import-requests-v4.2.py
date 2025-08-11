"""
Script Python per interrogare l'API Jira e generare un report delle issue assegnate all'utente corrente.

Funzionalità principali:
- Carica il token API Jira da un file .env per motivi di sicurezza.
- Esegue una query JQL per recuperare le issue assegnate all'utente con stati specifici (Da Gestire, In corso, Stand by Cliente, Stand by Interno).
- Raggruppa le issue per priorità (High, Medium, Low, Nessuna).
- Ordina le issue per data di scadenza e data di creazione.
- Genera due file di output:
  1. Un documento Word (.docx) con le issue formattate, in cui la chiave dell’issue è in grassetto.
  2. Un file di testo (.txt) con l’elenco delle issue.
- Gestisce eventuali errori di risposta dall’API.

Prerequisiti:
- Installare le librerie Python: requests, python-docx, python-dotenv
- Creare un file `.env` contenente la variabile JIRA_API_TOKEN con il token API di Jira.

Utilizzo:
- Modificare le variabili JIRA_URL e USERNAME con i propri dati.
- Eseguire lo script da riga di comando.
- Trovare i file `elenco_attivita.docx` e `elenco_attivita.txt` nella cartella di esecuzione.

Autore: Roberto Raimondi
"""

import requests
import sys
import os
import csv
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING

doc = Document()
load_dotenv()   

# Carica le variabili d'ambiente da .env se presente
VERSIONE    = "4.2"
JIRA_URL = os.getenv("JIRA_URL")               # Esempio: https://ejlog.atlassian.net
USERNAME = os.getenv("JIRA_USERNAME")          # Esempio: rraimondi@ferrettogroup.com
API_TOKEN = os.getenv("JIRA_API_TOKEN")

# JQL di ricerca
JQL = 'assignee = currentUser() AND status in ("Da Gestire", "In corso", "Stand by Cliente", "Stand by Interno") ORDER BY priority DESC, project, duedate ASC, created ASC'

# === PARAMETRI RICHIESTA ===
url = f"{JIRA_URL}/rest/api/3/search"
headers = {"Accept": "application/json"}
auth = (USERNAME, API_TOKEN)

all_issues = []
start_at = 0
max_results = 50  # limite massimo Jira Cloud

while True:
    params = {
        "jql": JQL,
        "fields": "summary,status,priority,created,duedate,project,key",
        "startAt": start_at,
        "maxResults": max_results
    }

    response = requests.get(url,
                            headers=headers, auth=auth, params=params)

    if response.status_code != 200:
        print(f"❌ Errore nella richiesta: {response.status_code} {response.text}")
        break

    data = response.json()

    issues = data.get("issues", [])
    if not issues:
        break

    all_issues.extend(issues)

    print(f"✅ Recuperati {len(issues)} ticket (totale finora: {len(all_issues)})")

    # Controlla se abbiamo preso tutto
    if start_at + max_results >= data.get("total", 0):
        break

    start_at += max_results

print(f"\n🎯 Recuperati in totale {len(all_issues)} ticket da Jira")

# Salvataggio CSV
csv_filename = "elenco_attivita.csv"
with open(csv_filename, mode="w", newline="", encoding="utf-8") as file:
    writer = csv.writer(file)
    # intestazioni CSV
    writer.writerow(["Key", "Summary", "Status", "Priority", "Created", "Due Date", "Project"])
    
    for issue in all_issues:
        fields = issue.get("fields", {})
        writer.writerow([
            issue.get("key", ""),
            fields.get("summary", ""),
            fields.get("status", {}).get("name", ""),
            fields.get("priority", {}).get("name", ""),
            fields.get("created", ""),
            fields.get("duedate", ""),
            fields.get("project", {}).get("key", "")
        ])

print(f"💾 File salvato: {csv_filename}")

# === CATEGORIZZAZIONE PER PRIORITÀ ===
priorities = {
    "Highest": [],
    "High": [],
    "Medium": [],
    "Low": [],
    "Lowest": [],
    "Nessuna": []
}

def parse_date(date_str):
    return datetime.strptime(date_str, "%Y-%m-%d") if date_str else None

def parse_created(date_str):
    return datetime.strptime(date_str, "%Y-%m-%d")

for issue in all_issues:
    fields = issue["fields"]
    key = issue["key"]
    project_name = fields["project"]["name"]
    title = fields["summary"]
    status = fields["status"]["name"]
    priority = fields.get("priority", {}).get("name", "Nessuna")
    duedate = fields.get("duedate")  # può essere None
    created = fields["created"][:10]

    title_clean = title
    if title.lower().startswith(project_name.lower()):
        title_clean = title[len(project_name):].lstrip(" -:–—")

    priorities.setdefault(priority, []).append({
        "key": key,
        "progetto": project_name,
        "titolo": title_clean,
        "stato": status,
        "scadenza": duedate,
        "creazione": created
    })

# === ORDINAMENTO E GENERAZIONE OUTPUT ===
output_lines = []

for prio_label in ["Highest", "High", "Medium", "Low", "Lowest", "Nessuna"]:
    blocco = priorities.get(prio_label, [])
    if not blocco:
        continue

    # Imposta margini pagina
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)      # margine superiore ridotto (0.5 cm)
        section.bottom_margin = Cm(0.5)   # margine inferiore ridotto (0.5 cm)
        section.left_margin = Cm(1)       # margine sinistro 1 cm
        section.right_margin = Cm(1)      # margine destro 1 cm

    # Definisci stile paragrafo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    # Per applicare Arial correttamente anche a caratteri asiatici ecc.
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Interlinea singola e nessuno spazio tra paragrafi dello stesso stile
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    # Sezione intestazione
    doc.add_paragraph("##############################")
    doc.add_paragraph(f"# {prio_label.upper()} PRIORITY")
    doc.add_paragraph("##############################\n")

    output_lines.append("##############################")
    output_lines.append(f"# {prio_label.upper()} PRIORITY")
    output_lines.append("##############################\n")

    # Ordinamento: prima per scadenza, poi per creazione
    blocco.sort(key=lambda x: (
        parse_date(x["scadenza"]) if x["scadenza"] else datetime.max,
        parse_created(x["creazione"])
    ))

    for item in blocco:
        scad = f", scad. {datetime.strptime(item['scadenza'], '%Y-%m-%d').strftime('%d-%m-%Y')}" if item["scadenza"] else ""
        line = f"{item['progetto']} - {item['titolo']} ({item['stato']}{scad})"

        # Word: key in grassetto
        p = doc.add_paragraph()
        run_key = p.add_run(f"{item['key']} ")
        run_key.bold = True
        p.add_run(line)

        # TXT: key inclusa
        output_lines.append(f"{item['key']} {line}")

    doc.add_paragraph("")
    output_lines.append("")

# === SALVA I FILE ===
doc.save("elenco_attivita.docx")

with open("elenco_attivita.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(output_lines))

print("✅ File 'elenco_attivita.docx' e 'elenco_attivita.txt' generati correttamente.")
