"""
Script Python per estrarre e documentare i dettagli di un progetto Jira.

Funzionalità principali:
- Recupera tramite API Jira le informazioni di un'issue principale (identificata da codice progetto, es. DNT-3).
- Estrae: titolo, riferimenti del cliente, informazioni sull'ambiente e l'elenco di tutte le issues collegate.
- Analizza e raccoglie tutti i commenti delle issues del progetto, ordinandoli per data.
- Genera un documento Word (.docx) strutturato con:
    - Intestazione progetto
    - Descrizione
    - Riferimenti cliente
    - Ambiente
    - Elenco issues
    - Commenti dettagliati
- Salva il report in un file `<CODICE_PROGETTO>_report.docx`.

Requisiti:
- Librerie Python: requests, python-docx
- API token Atlassian valido
- Permessi di accesso in lettura al progetto Jira
"""

import requests
import sys
from docx import Document
from docx.shared import Pt
from datetime import datetime

# === CONFIG ===
VERSIONE = "2.0"
JIRA_URL = "https://ejlog.atlassian.net"        # <-- Modifica con il tuo dominio
USERNAME = "rraimondi@ferrettogroup.com"        # <-- Tuo utente Atlassian
API_TOKEN = "ATATT3xFfGF0rFdHpQs25sHiT1VUn4-jmZl3nsdGWBB_hHAUMa1rePZsj758F_s7d0lWi7C0tCD4a-e8EOTQTvFZO_jtl3WYU0gG5tNv8bg2RkUDV24oYBapEB7-MQF7opXFmFREB1UCMDqRN_-ZZAkQL8xFceDjOowuNTx53CSyB1-HO1WJhUc=9EAB0E42"                  # <-- Generato da https://id.atlassian.com/manage/api-tokens

# === ID dei campi custom ===
CAMPO_RIFERIMENTI = "customfield_10146"
CAMPO_AMBIENTE = "environment"

def estrai_testo_rich_text(campo):
    if not campo or "content" not in campo:
        return ""
    testo = ""
    for blocco in campo["content"]:
        if blocco["type"] == "text":
            testo += blocco.get("text", "")
        elif blocco["type"] == "hardBreak":
            testo += "\n"
        else:
            # ricorsione per contenuti annidati
            testo += estrai_testo_rich_text(blocco)
        # Aggiungo una nuova riga dopo ogni blocco (es. paragrafo)
        if blocco["type"] in ("paragraph", "heading", "blockquote"):
            testo += "\n"
    return testo.strip()

def estrai_commenti_progetto(jira_url, auth, issues, codice_progetto):
    tutti_commenti = []
    for issue in issues:
        if not issue["key"].startswith(codice_progetto):
            continue
        issue_key = issue["key"]
        url_commenti = f"{jira_url}/rest/api/3/issue/{issue_key}/comment"
        resp = requests.get(url_commenti, auth=auth, headers={"Accept": "application/json"})
        if resp.status_code == 200:
            dati_commenti = resp.json().get("comments", [])
            for c in dati_commenti:
                data_commento = datetime.strptime(c["created"][:19], "%Y-%m-%dT%H:%M:%S")
                data_formattata = data_commento.strftime("%d %b %Y")  # '26 set 2023'
                autore = c["author"]["displayName"]

                # Qui elaboriamo il corpo come rich text
                body_field = c.get("body", {})
                testo = estrai_testo_rich_text(c.get("body", {}))

                tutti_commenti.append({
                    "issue_key": issue_key,
                    "created": data_commento,
                    "author": autore,
                    "body": testo
                })
    tutti_commenti.sort(key=lambda x: x["created"])
    return tutti_commenti

# === MAIN ===
if len(sys.argv) < 2:
    print("❌ Inserire il codice progetto (es: DNT-3)")
    sys.exit(1)

HEADERS = {"Accept": "application/json"}
AUTH = (USERNAME, API_TOKEN)
CODICE_PROGETTO = sys.argv[1].upper()
PROJECT_KEY = CODICE_PROGETTO.split("-")[0]

# Recupero dati issue principale (quella del codice completo esatto)
issue_url = f"{JIRA_URL}/rest/api/3/issue/{CODICE_PROGETTO}"
resp = requests.get(issue_url, headers=HEADERS, auth=AUTH)

if resp.status_code != 200:
    print(f"❌ Errore nel recupero dell'issue {CODICE_PROGETTO}")
    print(resp.text)
    sys.exit(1)

issue_data = resp.json()
summary = issue_data["fields"]["summary"]

riferimenti_raw = issue_data["fields"].get(CAMPO_RIFERIMENTI, None)
#riferimenti = estrai_testo_da_campo_rich_text(riferimenti_raw)
riferimenti = estrai_testo_rich_text(riferimenti_raw)

ambiente_raw = issue_data["fields"].get(CAMPO_AMBIENTE, None)
#ambiente = estrai_testo_da_campo_rich_text(ambiente_raw)
ambiente = estrai_testo_rich_text(ambiente_raw)

# Recupero tutte le issue del progetto
jql = f'project = "{PROJECT_KEY}" ORDER BY created ASC'
search_url = f"{JIRA_URL}/rest/api/3/search"
params = {
    "jql": jql,
    "fields": "summary",
    "maxResults": 1000
}

resp_all = requests.get(search_url, headers=HEADERS, params=params, auth=AUTH)
all_issues = resp_all.json().get("issues", [])

# Filtra solo le issue esatte del progetto (es. tutte DNT-3, DNT-4 scartate)
filtered_issues = [iss for iss in all_issues if iss["key"].startswith(CODICE_PROGETTO)]

# Estrai commenti solo da queste issue
commenti = estrai_commenti_progetto(JIRA_URL, AUTH, filtered_issues, CODICE_PROGETTO)

# Creazione documento Word
doc = Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)

# Intestazione
add_heading(f"Documento Progetto - {CODICE_PROGETTO}", level=0)

# Descrizione
add_heading("Descrizione")
add_paragraph(summary)

# Riferimenti
add_heading("Riferimenti delle persone del cliente")
add_paragraph(riferimenti)

# Ambiente
add_heading("Informazioni sull'Ambiente")
add_paragraph(ambiente)

# Elenco Issues
add_heading("Elenco Issues del Progetto")
for issue in filtered_issues:
    key = issue["key"]
    title = issue["fields"]["summary"]
    doc.add_paragraph(f"{key} - {title}", style="List Bullet")

doc.add_page_break()

# Commenti
add_heading("Commenti di tutte le issues del progetto", level=1)
for c in commenti:
    created_str = c["created"].strftime("%Y-%m-%d %H:%M")
#    doc.add_paragraph(f"{c['issue_key']} - {created_str} - {c['author']}")
    p = doc.add_paragraph()
    run = p.add_run(f"{c['issue_key']} - {created_str} - {c['author']}")
    run.bold = True
    doc.add_paragraph(c["body"])
    doc.add_paragraph("")

# Salvataggio
output_filename = f"{CODICE_PROGETTO}_report.docx"
doc.save(output_filename)

print(f"✅ Documento generato: {output_filename}")
